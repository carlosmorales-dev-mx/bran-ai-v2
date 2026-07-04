import csv
import tempfile
from io import BytesIO, StringIO
from pathlib import Path

from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

from google import genai
from google.genai import types

from core.config import GEMINI_API_KEY, GEMINI_MODEL


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".csv",
    ".txt",
    ".mp4",
    ".mov",
    ".jpg",
    ".jpeg",
    ".png",
}


def get_extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def detect_type(filename: str) -> str:
    ext = get_extension(filename)

    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".xlsx":
        return "xlsx"
    if ext == ".csv":
        return "csv"
    if ext == ".txt":
        return "text"
    if ext in [".jpg", ".jpeg", ".png"]:
        return "image"
    if ext in [".mp4", ".mov"]:
        return "video"

    return "file"


def parse_uploaded_file(
    filename: str,
    content: bytes,
    mime_type: str | None = None,
) -> str:
    ext = get_extension(filename)

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".txt":
        return parse_txt(content)

    if ext == ".csv":
        return parse_csv(content)

    if ext == ".xlsx":
        return parse_xlsx(content)

    if ext == ".docx":
        return parse_docx(content)

    if ext == ".pdf":
        return parse_pdf(content)

    if ext in [".jpg", ".jpeg", ".png"]:
        return parse_image_with_gemini(content, mime_type or guess_mime(ext))

    if ext in [".mp4", ".mov"]:
        return parse_video_with_gemini(
            content,
            mime_type or guess_mime(ext),
            filename,
        )

    raise ValueError(f"Unsupported file type: {ext}")


def parse_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="ignore").strip()


def parse_csv(content: bytes) -> str:
    text = content.decode("utf-8", errors="ignore")
    reader = csv.reader(StringIO(text))

    lines = []

    for row in reader:
        cleaned = [cell.strip() for cell in row if cell is not None]
        if cleaned:
            lines.append(" | ".join(cleaned))

    return "\n".join(lines).strip()


def parse_xlsx(content: bytes) -> str:
    wb = load_workbook(BytesIO(content), data_only=True)
    lines = []

    for sheet in wb.worksheets:
        lines.append(f"Hoja: {sheet.title}")

        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None]
            if values:
                lines.append(" | ".join(values))

    return "\n".join(lines).strip()


def parse_docx(content: bytes) -> str:
    doc = Document(BytesIO(content))
    lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]
            if values:
                lines.append(" | ".join(values))

    return "\n".join(lines).strip()


def parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    lines = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()

        if text:
            lines.append(f"Página {i + 1}")
            lines.append(text)

    return "\n\n".join(lines).strip()


def parse_image_with_gemini(content: bytes, mime_type: str) -> str:
    if not GEMINI_API_KEY:
        raise Exception("GEMINI_API_KEY no definida para procesar imagen")

    client = genai.Client(api_key=GEMINI_API_KEY)

    part = types.Part.from_bytes(
        data=content,
        mime_type=mime_type,
    )

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=[
            """
Extrae información de esta imagen para usarla como contexto privado de una conversación.

Reglas estrictas:
1. No respondas al usuario.
2. No hagas análisis general.
3. No hagas resumen interpretativo.
4. No agregues datos externos.
5. No inventes información.
6. No expliques qué es la imagen más allá de lo visible.
7. Extrae primero el texto visible tal como aparece.
8. Si hay elementos visuales importantes, descríbelos brevemente.
9. Mantén el resultado compacto.
10. Si hay texto cortado o ilegible, márcalo como parcial o ilegible.

Formato exacto:

TEXTO_VISIBLE:
- ...

ELEMENTOS_VISUALES:
- ...
""",
            part,
        ],
    )

    return (response.text or "").strip()


def parse_video_with_gemini(content: bytes, mime_type: str, filename: str) -> str:
    if not GEMINI_API_KEY:
        raise Exception("GEMINI_API_KEY no definida para procesar video")

    client = genai.Client(api_key=GEMINI_API_KEY)

    suffix = get_extension(filename)

    with tempfile.NamedTemporaryFile(delete=True, suffix=suffix) as tmp:
        tmp.write(content)
        tmp.flush()

        uploaded = client.files.upload(file=tmp.name)

        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=[
                uploaded,
                """
Extrae información de este video para usarla como contexto privado de una conversación.

Reglas estrictas:
1. No respondas al usuario.
2. No hagas análisis general.
3. No agregues datos externos.
4. No inventes información.
5. Transcribe únicamente lo audible si existe.
6. Extrae texto visible si aparece en pantalla.
7. Describe brevemente escenas relevantes.
8. Mantén el resultado compacto.

Formato exacto:

TRANSCRIPCION:
- ...

TEXTO_VISIBLE:
- ...

ESCENAS_RELEVANTES:
- ...
""",
            ],
        )

    return (response.text or "").strip()


def guess_mime(ext: str) -> str:
    if ext in [".jpg", ".jpeg"]:
        return "image/jpeg"
    if ext == ".png":
        return "image/png"
    if ext == ".mp4":
        return "video/mp4"
    if ext == ".mov":
        return "video/quicktime"

    return "application/octet-stream"